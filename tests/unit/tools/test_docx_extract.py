"""TDD tests for Task 19 — .docx extraction in tools.knowledge_ingest.

Covers:
- _extract_docx with python-docx present: returns paragraph + table text
- _extract_docx with python-docx absent (simulated): returns None, no raise
- _extract_docx with corrupt bytes (.docx name but garbage content): returns None
- kb_ingest over a dir containing a .docx:
    - lib present → docx ingested (chunks land in fake registry)
    - lib absent  → counted as skipped_office, not ingested
"""
from __future__ import annotations

import asyncio
import importlib
import sys
from pathlib import Path
from unittest.mock import patch

import pytest


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _run(coro):
    return asyncio.run(coro)


class FakeRegistry:
    """Minimal registry double (mirrors the one in test_knowledge_ingest.py)."""

    def __init__(self):
        self.ingested_chunks: list[dict] = []
        self.hashes: dict[str, str] = {}

    async def kb_ingest_chunk(self, *, user_id, collection, source_path,
                               source_hash, chunk_idx, content, mime, created_at):
        self.ingested_chunks.append({
            "source_path": source_path,
            "content": content,
        })
        self.hashes[source_path] = source_hash

    async def kb_remove(self, *, user_id, collection, source=None):
        if source and source in self.hashes:
            del self.hashes[source]
        return 1

    async def kb_source_hash(self, *, user_id, collection, source_path):
        return self.hashes.get(source_path)


def _make_docx(path: Path, paragraphs: list[str], table_rows: list[list[str]] | None = None) -> None:
    """Write a real .docx to *path* using python-docx."""
    import docx as _docx
    doc = _docx.Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    if table_rows:
        n_cols = max(len(r) for r in table_rows)
        tbl = doc.add_table(rows=len(table_rows), cols=n_cols)
        for r_idx, row in enumerate(table_rows):
            for c_idx, cell_text in enumerate(row):
                tbl.rows[r_idx].cells[c_idx].text = cell_text
    doc.save(str(path))


def _run_kb_ingest(path_str: str, tmp_path: Path, fake: FakeRegistry, **kwargs):
    """Run kb_ingest with patched registry + confinement root."""
    import tools.knowledge_ingest as ki_mod

    async def _go():
        return await ki_mod.kb_ingest(path_str, user_id="u1", session_id="s1", **kwargs)

    with patch.object(ki_mod, "_resolve_confinement_root", return_value=tmp_path), \
         patch("modules.memory.registry.kb_ingest_chunk", new=fake.kb_ingest_chunk), \
         patch("modules.memory.registry.kb_remove", new=fake.kb_remove), \
         patch("modules.memory.registry.kb_source_hash", new=fake.kb_source_hash):
        return _run(_go())


# ---------------------------------------------------------------------------
# _extract_docx — python-docx present
# ---------------------------------------------------------------------------


class TestExtractDocxPresent:
    def test_returns_paragraph_text(self, tmp_path):
        from tools.knowledge_ingest import _extract_docx

        docx_path = tmp_path / "test.docx"
        _make_docx(docx_path, ["Hello world", "Second paragraph"])

        result = _extract_docx(docx_path)

        assert result is not None
        assert "Hello world" in result
        assert "Second paragraph" in result

    def test_multiple_paragraphs_joined_by_newline(self, tmp_path):
        from tools.knowledge_ingest import _extract_docx

        docx_path = tmp_path / "multi.docx"
        _make_docx(docx_path, ["Line one", "Line two", "Line three"])

        result = _extract_docx(docx_path)

        assert result is not None
        lines = result.split("\n")
        texts = [l.strip() for l in lines if l.strip()]
        assert "Line one" in texts
        assert "Line two" in texts
        assert "Line three" in texts

    def test_table_cell_text_included(self, tmp_path):
        from tools.knowledge_ingest import _extract_docx

        docx_path = tmp_path / "table.docx"
        _make_docx(docx_path, ["Intro para"], table_rows=[["Cell A", "Cell B"], ["Cell C", "Cell D"]])

        result = _extract_docx(docx_path)

        assert result is not None
        assert "Intro para" in result
        assert "Cell A" in result
        assert "Cell D" in result

    def test_returns_string_type(self, tmp_path):
        from tools.knowledge_ingest import _extract_docx

        docx_path = tmp_path / "typed.docx"
        _make_docx(docx_path, ["Some content"])
        result = _extract_docx(docx_path)
        assert isinstance(result, str)


# ---------------------------------------------------------------------------
# _extract_docx — python-docx absent (simulated)
# ---------------------------------------------------------------------------


class TestExtractDocxAbsent:
    def test_returns_none_when_import_fails(self, tmp_path):
        from tools.knowledge_ingest import _extract_docx

        docx_path = tmp_path / "any.docx"
        _make_docx(docx_path, ["content"])

        # Simulate python-docx not installed by making the import raise
        real_import = __builtins__.__import__ if hasattr(__builtins__, "__import__") else __import__

        def _patched_import(name, *args, **kwargs):
            if name == "docx":
                raise ImportError("No module named 'docx'")
            return real_import(name, *args, **kwargs)

        with patch("builtins.__import__", side_effect=_patched_import):
            result = _extract_docx(docx_path)

        assert result is None

    def test_no_exception_raised_when_import_fails(self, tmp_path):
        """Absent lib must never propagate an exception."""
        from tools.knowledge_ingest import _extract_docx

        docx_path = tmp_path / "any.docx"
        _make_docx(docx_path, ["content"])

        real_import = __builtins__.__import__ if hasattr(__builtins__, "__import__") else __import__

        def _patched_import(name, *args, **kwargs):
            if name == "docx":
                raise ImportError("No module named 'docx'")
            return real_import(name, *args, **kwargs)

        with patch("builtins.__import__", side_effect=_patched_import):
            # Must not raise
            result = _extract_docx(docx_path)

        assert result is None


# ---------------------------------------------------------------------------
# _extract_docx — corrupt file
# ---------------------------------------------------------------------------


class TestExtractDocxCorrupt:
    def test_corrupt_docx_returns_none(self, tmp_path):
        from tools.knowledge_ingest import _extract_docx

        corrupt = tmp_path / "corrupt.docx"
        corrupt.write_bytes(b"\x00\x01\x02\x03garbage bytes not a valid zip")

        result = _extract_docx(corrupt)
        assert result is None

    def test_corrupt_docx_no_exception(self, tmp_path):
        """A corrupt .docx must never propagate an exception."""
        from tools.knowledge_ingest import _extract_docx

        corrupt = tmp_path / "corrupt.docx"
        corrupt.write_bytes(b"PK\x03\x04notreallyadocx")

        # Should not raise
        result = _extract_docx(corrupt)
        assert result is None


# ---------------------------------------------------------------------------
# _extract_text routing for .docx
# ---------------------------------------------------------------------------


class TestExtractTextDocxRouting:
    def test_docx_routed_to_docx_extractor(self, tmp_path):
        """_extract_text for a .docx returns (text, None) when lib present."""
        from tools.knowledge_ingest import _extract_text

        docx_path = tmp_path / "routed.docx"
        _make_docx(docx_path, ["Routing test paragraph"])

        text, reason = _run(_extract_text(docx_path))

        assert reason is None
        assert text is not None
        assert "Routing test paragraph" in text

    def test_docx_skipped_with_note_when_lib_absent(self, tmp_path):
        """When _extract_docx returns None (lib absent), reason starts with 'office-skip'."""
        from tools.knowledge_ingest import _extract_text

        docx_path = tmp_path / "nolib.docx"
        _make_docx(docx_path, ["content"])

        with patch("tools.knowledge_ingest._extract_docx", return_value=None):
            text, reason = _run(_extract_text(docx_path))

        assert text is None
        assert reason is not None
        assert reason.startswith("office-skip")

    def test_other_office_still_skipped(self, tmp_path):
        """Non-.docx office formats (.doc, .xlsx) still get office-skip."""
        from tools.knowledge_ingest import _extract_text

        for ext in (".doc", ".xlsx", ".pptx"):
            f = tmp_path / f"file{ext}"
            f.write_bytes(b"fake content")
            text, reason = _run(_extract_text(f))
            assert text is None
            assert reason is not None
            assert reason.startswith("office-skip"), f"{ext} should be office-skip"


# ---------------------------------------------------------------------------
# kb_ingest integration — .docx with lib present
# ---------------------------------------------------------------------------


class TestKbIngestDocxPresent:
    def test_docx_file_ingested(self, tmp_path):
        """A .docx in a directory is ingested when python-docx is present."""
        docx_path = tmp_path / "report.docx"
        _make_docx(docx_path, ["Important finding one", "Important finding two"])

        fake = FakeRegistry()
        result = _run_kb_ingest(str(tmp_path), tmp_path, fake)

        assert result["ingested"] == 1
        assert result["n_chunks"] >= 1
        assert result.get("skipped_office", 0) == 0
        # Confirm the chunk content came from the docx paragraphs
        all_content = " ".join(c["content"] for c in fake.ingested_chunks)
        assert "Important finding" in all_content

    def test_docx_single_file_ingested(self, tmp_path):
        """Single-file .docx ingestion path works."""
        docx_path = tmp_path / "single.docx"
        _make_docx(docx_path, ["Single file paragraph"])

        fake = FakeRegistry()
        result = _run_kb_ingest(str(docx_path), tmp_path, fake)

        assert result["ingested"] == 1
        assert result.get("skipped_office", 0) == 0


# ---------------------------------------------------------------------------
# kb_ingest integration — .docx with lib absent (simulated)
# ---------------------------------------------------------------------------


class TestKbIngestDocxAbsent:
    def test_docx_counted_as_skipped_office_when_lib_absent(self, tmp_path):
        """When python-docx is absent, .docx counts as skipped_office, not ingested."""
        docx_path = tmp_path / "nopkg.docx"
        _make_docx(docx_path, ["Some text"])

        fake = FakeRegistry()

        # Simulate absent library by patching _extract_docx to return None
        with patch("tools.knowledge_ingest._extract_docx", return_value=None):
            result = _run_kb_ingest(str(tmp_path), tmp_path, fake)

        assert result["ingested"] == 0
        assert result.get("skipped_office", 0) >= 1
        assert len(fake.ingested_chunks) == 0

    def test_docx_absent_lib_no_error_key(self, tmp_path):
        """Absent lib produces skipped_office, not an 'error' key in result."""
        docx_path = tmp_path / "nopkg2.docx"
        _make_docx(docx_path, ["content"])

        fake = FakeRegistry()

        with patch("tools.knowledge_ingest._extract_docx", return_value=None):
            result = _run_kb_ingest(str(docx_path), tmp_path, fake)

        assert "error" not in result
        assert result.get("skipped_office", 0) >= 1
