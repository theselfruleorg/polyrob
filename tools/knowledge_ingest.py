"""KB ingestion engine + `knowledge` agent tool (Task 6).

Adds file→KB ingestion and a registerable ``knowledge`` tool on top of the
Task-5 registry routers.  All KB I/O goes through ``modules.memory.registry``
(no direct provider access here).

``from __future__ import annotations`` is SAFE in this file — the param models are
explicit (passed to ``BaseTool.action``), so the Registry never needs to introspect
stringized first-arg annotations here.
"""
from __future__ import annotations

import asyncio
import hashlib
import logging
import os
import re
import subprocess
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from pydantic import BaseModel, ConfigDict, Field

from tools.base_tool import BaseTool
from tools.controller.types import ActionResult
from agents.task.agent.core.secret_guard import (
    is_secret_path,
    is_binary_file,
    estimate_tokens_rough,
)
from core.path_safety import is_within_root

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Env-driven knobs (read inline, not in constants.py)
# ---------------------------------------------------------------------------


def _int_env(name: str, default: int) -> int:
    try:
        return int(os.getenv(name, str(default)))
    except (TypeError, ValueError):
        return default


def _kb_max_files() -> int:
    return _int_env("KB_MAX_FILES", 2000)


def _kb_max_bytes() -> int:
    return _int_env("KB_MAX_BYTES", 25 * 1024 * 1024)


def _kb_chunk_tokens() -> int:
    return _int_env("KB_CHUNK_TOKENS", 800)


def _kb_chunk_overlap() -> int:
    return _int_env("KB_CHUNK_OVERLAP", 100)


# ---------------------------------------------------------------------------
# Office / unsupported extension skip list
# ---------------------------------------------------------------------------

_OFFICE_EXTENSIONS: frozenset[str] = frozenset({
    ".doc", ".odt",
    ".xlsx", ".xls", ".ods",
    ".pptx", ".ppt",
    ".epub", ".rtf",
})

# Extensions handled by dedicated extractors (not in _OFFICE_EXTENSIONS skip list)
_DOCX_EXTENSION = ".docx"

_TEXT_EXTENSIONS: frozenset[str] = frozenset({
    ".txt", ".md", ".rst", ".csv",
    ".py", ".js", ".ts", ".tsx", ".jsx",
    ".sh", ".bash", ".zsh",
    ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg",
    ".html", ".xml", ".sql",
    ".go", ".rs", ".c", ".cpp", ".h", ".hpp",
    ".java", ".kt", ".swift", ".rb", ".php",
    ".r", ".scala", ".clj",
})


# ---------------------------------------------------------------------------
# Pure engine helpers
# ---------------------------------------------------------------------------


def _iter_files(
    root: Path,
    *,
    recursive: bool = True,
    globs: Optional[List[str]] = None,
    max_files: int = 2000,
    max_bytes: int = 25 * 1024 * 1024,
) -> Tuple[List[Path], Dict[str, int]]:
    """Bounded, safe walk of *root*.

    Returns ``(files, skipped)`` where ``skipped`` maps reason→count.
    Prefers ``rg --files`` for gitignore-aware traversal; falls back to
    ``os.walk`` skipping hidden dirs and ``__pycache__``.

    Hard-skips:
    - ``is_secret_path`` (credentials / env files)
    - ``is_binary_file`` (compiled binaries, images, etc.)
    - Files beyond ``max_files`` / ``max_bytes`` cumulative size limits.
    """
    root = root.resolve()
    skipped: Dict[str, int] = {
        "secret": 0,
        "binary": 0,
        "max_files": 0,
        "max_bytes": 0,
    }

    # Collect candidate paths ------------------------------------------------
    candidates: List[Path] = []

    if recursive:
        # Try rg --files for .gitignore awareness
        rg_paths = _rg_files(root)
        if rg_paths is not None:
            candidates = rg_paths
        else:
            # Fallback: os.walk, skip hidden dirs + __pycache__
            for dirpath, dirnames, filenames in os.walk(str(root)):
                # Prune hidden dirs and __pycache__ in-place
                dirnames[:] = [
                    d for d in dirnames
                    if not d.startswith(".") and d != "__pycache__"
                ]
                for fname in filenames:
                    candidates.append(Path(dirpath) / fname)
    else:
        # Non-recursive: direct children only
        try:
            candidates = [p for p in root.iterdir() if p.is_file()]
        except OSError:
            pass

    # Apply glob filters if requested ----------------------------------------
    if globs:
        import fnmatch as _fnmatch
        filtered = []
        for p in candidates:
            rel = str(p.relative_to(root)) if p.is_absolute() else str(p)
            if any(_fnmatch.fnmatch(rel, g) or _fnmatch.fnmatch(p.name, g) for g in globs):
                filtered.append(p)
        candidates = filtered

    # Filter: secret / binary; enforce limits --------------------------------
    collected: List[Path] = []
    total_bytes = 0

    for path in candidates:
        if not path.is_file():
            continue

        # Secret hard-skip
        if is_secret_path(path, root=root):
            skipped["secret"] += 1
            continue

        # Binary hard-skip (don't hard-skip .pdf — extractable)
        if is_binary_file(path):
            skipped["binary"] += 1
            continue

        # File count cap
        if len(collected) >= max_files:
            skipped["max_files"] += 1
            continue

        # Byte cap (skip oversized files; cumulative cap too)
        try:
            fsize = path.stat().st_size
        except OSError:
            continue
        if total_bytes + fsize > max_bytes:
            skipped["max_bytes"] += 1
            continue

        collected.append(path)
        total_bytes += fsize

    return collected, skipped


def _rg_files(root: Path) -> Optional[List[Path]]:
    """Run ``rg --files <root>`` and return the paths, or None if unavailable."""
    try:
        result = subprocess.run(
            ["rg", "--files", str(root)],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode != 0:
            return None
        paths = []
        for line in result.stdout.splitlines():
            line = line.strip()
            if line:
                p = Path(line)
                if not p.is_absolute():
                    p = root / p
                paths.append(p)
        return paths
    except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
        return None


async def _extract_text(path: Path) -> Tuple[Optional[str], Optional[str]]:
    """Extract text from *path*.

    Returns ``(text, skip_reason)`` — exactly one of the two is ``None``.

    - ``.pdf`` → direct ``await PdfExtractionMixin._process_pdf(content)``
      (sync ``_extract_pdf_direct`` fallback only if the mixin raises / imports fail)
    - text / code extensions → ``path.read_text(errors='replace')``
    - office formats → ``(None, "office-skip")`` — no parser yet (Task 19)
    """
    suffix = path.suffix.lower()

    # .docx — extract via python-docx (fail-open: if lib absent → skip-with-note)
    if suffix == _DOCX_EXTENSION:
        text = _extract_docx(path)
        if text is None:
            return None, f"office-skip:{suffix}"
        return text, None

    # Other office formats — no parser available
    if suffix in _OFFICE_EXTENSIONS:
        return None, f"office-skip:{suffix}"

    # PDF extraction (async — directly await the mixin, no sync-from-async bridge)
    if suffix == ".pdf":
        try:
            content_bytes = path.read_bytes()
            text = await _extract_pdf_text(content_bytes)
            if text:
                return text, None
            return None, "pdf-empty"
        except Exception as e:
            logger.debug("PDF extraction failed for %s: %s", path, e)
            return None, f"pdf-error:{type(e).__name__}"

    # Plain text / code
    try:
        text = path.read_text(errors="replace")
        return text, None
    except OSError as e:
        return None, f"read-error:{type(e).__name__}"


async def _extract_pdf_text(content_bytes: bytes) -> str:
    """Async adapter over ``PdfExtractionMixin._process_pdf``.

    Directly ``await``s the (async) mixin from within the running loop — no
    ``run_coroutine_sync`` bridge.  Falls back to the SYNC ``_extract_pdf_direct``
    only if the mixin raises or its import fails.
    """
    try:
        from tools.filesystem_pdf import PdfExtractionMixin
        mixin = PdfExtractionMixin()
        # _process_pdf reads self.logger / self.container — provide minimal stand-ins.
        mixin.logger = logging.getLogger("knowledge_ingest.pdf")  # type: ignore[attr-defined]
        mixin.container = None  # type: ignore[attr-defined]
        result = await mixin._process_pdf(content_bytes)
        return result.get("content", "")
    except Exception as e:
        logger.debug("PDF mixin extraction failed, using direct fallback: %s", e)
        return _extract_pdf_direct(content_bytes)


def _extract_pdf_direct(content_bytes: bytes) -> str:
    """SYNC-only fallback PDF extraction using pypdf directly.

    Used only when ``_process_pdf`` raises or its import fails.  No event-loop
    detection, no ``run_coroutine_sync``.
    """
    try:
        import pypdf  # type: ignore
        from io import BytesIO
        reader = pypdf.PdfReader(BytesIO(content_bytes), strict=False)
        texts = []
        for page in reader.pages:
            try:
                t = page.extract_text()
                if t:
                    texts.append(t)
            except Exception:
                pass
        return "\n\n".join(texts)
    except Exception as e:
        logger.debug("Direct PDF extraction failed: %s", e)
        return ""


def _extract_docx(path: Path) -> Optional[str]:
    """Extract plain text from a ``.docx`` file via ``python-docx``.

    Lazy import — if ``python-docx`` is not installed, returns ``None`` so
    the caller skips the file with-note (fail-open, never an error).

    Paragraph text is joined with ``\\n``; table cell text is appended after
    paragraphs so tables aren't silently dropped.

    Any extraction error (corrupt file, unexpected docx layout) also returns
    ``None`` (fail-open, logged at DEBUG).
    """
    try:
        import docx as _docx  # python-docx; import name is 'docx'
    except ImportError:
        logger.debug("python-docx not installed; skipping %s", path)
        return None

    try:
        doc = _docx.Document(str(path))
        parts: list[str] = []

        # Body paragraphs
        for para in doc.paragraphs:
            text = para.text
            if text.strip():
                parts.append(text)

        # Table cells (flattened, row-major)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        parts.append(cell_text)

        return "\n".join(parts) if parts else ""
    except Exception as e:
        logger.debug("docx extraction failed for %s: %s", path, e)
        return None


# ---------------------------------------------------------------------------
# Structure-first chunker
# ---------------------------------------------------------------------------

# Markdown heading pattern
_MD_HEADING_RE = re.compile(r"^#{1,6} .+", re.MULTILINE)
# Top-level Python def/class
_PYDEF_RE = re.compile(r"^(?:def |class )\S", re.MULTILINE)


def _chunk(
    text: str,
    *,
    target: Optional[int] = None,
    overlap: Optional[int] = None,
) -> List[str]:
    """Structure-first text chunker.

    Splits on: markdown headings → blank-line paragraphs → top-level
    ``def``/``class`` → raw size.  Chunks are sized by ``estimate_tokens_rough``
    to ~``target`` tokens with ``overlap`` token carry.
    """
    target = target if target is not None else _kb_chunk_tokens()
    overlap = overlap if overlap is not None else _kb_chunk_overlap()

    if not text or not text.strip():
        return []

    # 1. Try markdown heading split
    parts = _split_on_headings(text)
    if len(parts) <= 1:
        # 2. Try blank-line paragraph split
        parts = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if len(parts) <= 1:
        # 3. Try top-level def/class split (Python)
        parts = _split_on_defs(text)
    if len(parts) <= 1:
        # 4. Raw line split
        parts = text.splitlines()

    # Merge small parts up to target size; carry overlap into next chunk
    return _merge_to_target(parts, target=target, overlap=overlap)


def _split_on_headings(text: str) -> List[str]:
    """Split *text* at markdown ``#`` headings (keep heading in chunk)."""
    positions = [m.start() for m in _MD_HEADING_RE.finditer(text)]
    if not positions:
        return [text]
    parts = []
    positions.append(len(text))
    for i in range(len(positions) - 1):
        chunk = text[positions[i]:positions[i + 1]].strip()
        if chunk:
            parts.append(chunk)
    # Prepend any text before the first heading
    if positions[0] > 0:
        preamble = text[:positions[0]].strip()
        if preamble:
            parts.insert(0, preamble)
    return parts if parts else [text]


def _split_on_defs(text: str) -> List[str]:
    """Split on top-level Python ``def``/``class`` lines."""
    positions = [m.start() for m in _PYDEF_RE.finditer(text)]
    if not positions:
        return [text]
    parts = []
    positions.append(len(text))
    if positions[0] > 0:
        preamble = text[:positions[0]].strip()
        if preamble:
            parts.append(preamble)
    for i in range(len(positions) - 1):
        chunk = text[positions[i]:positions[i + 1]].strip()
        if chunk:
            parts.append(chunk)
    return parts if parts else [text]


def _merge_to_target(parts: List[str], *, target: int, overlap: int) -> List[str]:
    """Merge *parts* into chunks of at most *target* tokens, carrying *overlap*.

    If a single *part* exceeds *target* tokens, it is further split on word
    boundaries so no output chunk grows unboundedly.
    """
    chunks: List[str] = []
    current_lines: List[str] = []
    current_tokens = 0
    overlap_carry = ""

    def _flush():
        nonlocal current_lines, current_tokens, overlap_carry
        if not current_lines:
            return
        chunk_text = overlap_carry + ("\n\n" if overlap_carry else "") + "\n\n".join(current_lines)
        chunks.append(chunk_text.strip())
        overlap_carry = _tail_tokens(chunk_text, overlap)
        current_lines = []
        current_tokens = estimate_tokens_rough(overlap_carry)

    for part in parts:
        part_tokens = estimate_tokens_rough(part)

        # If a single part is too large, split it on word boundaries first
        if part_tokens > target:
            # Flush anything accumulated so far
            _flush()
            # Split oversized part into word-boundary slices
            for sub in _split_by_tokens(part, target):
                sub_tokens = estimate_tokens_rough(sub)
                if current_tokens + sub_tokens > target and current_lines:
                    _flush()
                current_lines.append(sub)
                current_tokens += sub_tokens
            continue

        if current_tokens + part_tokens > target and current_lines:
            _flush()

        current_lines.append(part)
        current_tokens += part_tokens

    # Final flush
    _flush()

    return [c for c in chunks if c]


def _split_by_tokens(text: str, target: int) -> List[str]:
    """Split *text* into slices of at most *target* tokens (word-boundary split)."""
    words = text.split()
    slices: List[str] = []
    current_words: List[str] = []
    current_chars = 0
    target_chars = target * 4  # estimate_tokens_rough uses // 4

    for word in words:
        wlen = len(word) + 1  # +1 for space
        if current_chars + wlen > target_chars and current_words:
            slices.append(" ".join(current_words))
            current_words = [word]
            current_chars = wlen
        else:
            current_words.append(word)
            current_chars += wlen

    if current_words:
        slices.append(" ".join(current_words))

    return slices if slices else [text]


def _tail_tokens(text: str, n_tokens: int) -> str:
    """Return the last ~*n_tokens* token-equivalents of *text* (char-based approx)."""
    if n_tokens <= 0:
        return ""
    tail_chars = n_tokens * 4  # estimate_tokens_rough uses // 4
    return text[-tail_chars:].lstrip()


# ---------------------------------------------------------------------------
# Async ingestion entry point
# ---------------------------------------------------------------------------


async def kb_ingest(
    path: str,
    collection: str = "default",
    recursive: bool = True,
    globs: Optional[List[str]] = None,
    *,
    user_id: str,
    session_id: str,
    source_name: Optional[str] = None,
) -> Dict[str, Any]:
    """Ingest a file or directory into the KB.

    Returns counts: ``{ingested, skipped_secret, skipped_binary, unchanged,
    n_chunks, skipped_office}``.

    Path confinement:
    - Server sessions: confined to ``pm().get_workspace_dir(session_id, user_id)``
    - CLI (workspace_is_project_root): confined to ``Path.cwd()``

    ``source_name`` (optional): a STABLE logical identity to store as the chunk's
    ``source_path`` / dedup key instead of the volatile on-disk path. Used by the
    HTTP upload path so a re-upload of the same file dedups by its original filename
    rather than a fresh ``mkstemp`` path. Only honored for single-file ingestion
    (a directory walk has many sources, so ``source_name`` is ignored there). When
    ``None`` the behavior is byte-identical to before (on-disk path is the source).
    """
    from modules.memory.registry import (
        kb_ingest_chunk as _kb_ingest_chunk,
        kb_remove as _kb_remove,
        kb_source_hash as _kb_source_hash,
    )

    target_path = Path(path)

    # ------------------------------------------------------------------
    # Path confinement
    # ------------------------------------------------------------------
    allowed_root = _resolve_confinement_root(session_id, user_id)
    resolved = target_path.resolve() if target_path.is_absolute() else (allowed_root / path).resolve()

    if not is_within_root(str(resolved), str(allowed_root)):
        return {
            "error": f"Path '{path}' is outside the allowed workspace root '{allowed_root}'",
            "ingested": 0,
            "skipped_secret": 0,
            "skipped_binary": 0,
            "unchanged": 0,
            "n_chunks": 0,
            "skipped_office": 0,
            "skipped_too_large": 0,
            "failed": 0,
        }

    # ------------------------------------------------------------------
    # Walk files
    # ------------------------------------------------------------------
    if resolved.is_file():
        files = [resolved]
        walk_skipped: Dict[str, int] = {
            "secret": 0, "binary": 0, "max_files": 0, "max_bytes": 0, "too_large": 0,
        }
        # Still apply secret/binary guards for single-file ingestion
        if is_secret_path(resolved, root=allowed_root):
            walk_skipped["secret"] += 1
            files = []
        elif is_binary_file(resolved):
            walk_skipped["binary"] += 1
            files = []
        else:
            # Byte cap also applies to a single file — never read an unbounded file
            # into memory (the directory walk caps cumulative size; this caps the
            # one-file path, which would otherwise OOM on a multi-GB file).
            try:
                if resolved.stat().st_size > _kb_max_bytes():
                    walk_skipped["too_large"] += 1
                    files = []
            except OSError:
                files = []
    elif resolved.is_dir():
        files, walk_skipped = _iter_files(
            resolved,
            recursive=recursive,
            globs=globs,
            max_files=_kb_max_files(),
            max_bytes=_kb_max_bytes(),
        )
    else:
        return {
            "error": f"Path '{path}' does not exist or is not accessible.",
            "ingested": 0,
            "skipped_secret": 0,
            "skipped_binary": 0,
            "unchanged": 0,
            "n_chunks": 0,
            "skipped_office": 0,
            "skipped_too_large": 0,
            "failed": 0,
        }

    # ------------------------------------------------------------------
    # Ingest each file
    # ------------------------------------------------------------------
    counts: Dict[str, int] = {
        "ingested": 0,
        "skipped_secret": walk_skipped.get("secret", 0),
        "skipped_binary": walk_skipped.get("binary", 0),
        "unchanged": 0,
        "n_chunks": 0,
        "skipped_office": 0,
        # oversized single file ("too_large") + oversized files in a dir walk ("max_bytes")
        "skipped_too_large": walk_skipped.get("too_large", 0) + walk_skipped.get("max_bytes", 0),
        "failed": 0,
    }

    target_tokens = _kb_chunk_tokens()
    overlap_tokens = _kb_chunk_overlap()

    # A stable source_name only makes sense for single-file ingestion (a directory
    # walk has many distinct sources). Ignore it for multi-file walks.
    use_source_name = source_name if (source_name and len(files) == 1) else None

    for fpath in files:
        source_path_str = use_source_name or str(fpath)

        # Read + hash off the event loop — reading a large file and hashing it is
        # synchronous CPU/IO that would otherwise stall every other session during a
        # bulk ingest.
        try:
            file_hash = await asyncio.to_thread(_read_and_hash, fpath)
        except OSError as e:
            logger.warning("Cannot read %s: %s", fpath, e)
            continue

        # Check existing hash via registry
        existing_hash = await _kb_source_hash(
            user_id=user_id,
            collection=collection,
            source_path=source_path_str,
        )

        if existing_hash == file_hash:
            counts["unchanged"] += 1
            continue

        # If changed (or new): remove old chunks, re-ingest
        if existing_hash is not None:
            await _kb_remove(
                user_id=user_id,
                collection=collection,
                source=source_path_str,
            )

        # Extract text
        text, skip_reason = await _extract_text(fpath)
        if text is None:
            if skip_reason and skip_reason.startswith("office-skip"):
                counts["skipped_office"] += 1
            else:
                logger.debug("Skipping %s: %s", fpath, skip_reason)
            continue

        # Chunk (CPU-bound) off the event loop too.
        chunks = await asyncio.to_thread(
            _chunk, text, target=target_tokens, overlap=overlap_tokens
        )
        if not chunks:
            continue

        # Ingest each chunk
        mime = "application/pdf" if fpath.suffix.lower() == ".pdf" else "text/plain"
        created_at = datetime.now(timezone.utc).isoformat()

        file_ok = True
        for idx, chunk_text in enumerate(chunks):
            ok = await _kb_ingest_chunk(
                user_id=user_id,
                collection=collection,
                source_path=source_path_str,
                source_hash=file_hash,
                chunk_idx=idx,
                content=chunk_text,
                mime=mime,
                created_at=created_at,
            )
            # A provider that returns False signals a real write failure (None from a
            # legacy/no-op provider is treated as success — unchanged behavior).
            if ok is False:
                file_ok = False
                break

        if not file_ok:
            # Don't leave the file half-ingested with a current source hash (a re-run
            # would skip it as "unchanged" and never recover the dropped chunks).
            # Remove the partial source so the next ingest retries from scratch.
            try:
                await _kb_remove(
                    user_id=user_id,
                    collection=collection,
                    source=source_path_str,
                )
            except Exception as e:
                logger.debug("partial-ingest cleanup failed for %s: %s", fpath, e)
            counts["failed"] += 1
            continue

        counts["ingested"] += 1
        counts["n_chunks"] += len(chunks)

    return counts


def _read_and_hash(path: Path) -> str:
    """Read file bytes and return their sha256 hex digest (runs in a worker thread)."""
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _resolve_confinement_root(session_id: str, user_id: str) -> Path:
    """Return the allowed root for path confinement.

    Server: ``pm().get_workspace_dir(session_id, user_id)``
    CLI / local mode: ``Path.cwd()``.

    Fail-CLOSED on the server: if the workspace root can't be resolved (``pm()``
    raises), we must NOT silently widen confinement to the process CWD (e.g.
    ``/opt/rob`` — app code/config). Only local mode (single-user) accepts the
    CWD fallback; on the server the error propagates so the caller refuses.
    """
    try:
        from agents.task.path import pm
        path_manager = pm()
        workspace = path_manager.get_workspace_dir(session_id, user_id)
        return workspace.resolve()
    except Exception:
        from agents.task.constants import local_mode_enabled
        if local_mode_enabled():
            return Path.cwd().resolve()
        raise


# ---------------------------------------------------------------------------
# Param models
# ---------------------------------------------------------------------------


class KbIngestParams(BaseModel):
    model_config = ConfigDict(extra="forbid")
    path: str = Field(..., description="File or directory path to ingest into the KB.")
    collection: str = Field("default", description="KB collection name.")
    recursive: bool = Field(True, description="Recurse into subdirectories.")
    globs: Optional[List[str]] = Field(
        None, description="Optional glob patterns to restrict which files are included."
    )


class KbSearchParams(BaseModel):
    model_config = ConfigDict(extra="forbid")
    query: str = Field(..., description="Search query.")
    collection: str = Field("default", description="KB collection to search.")
    limit: int = Field(8, ge=1, le=50, description="Maximum number of results.")


class KbListParams(BaseModel):
    model_config = ConfigDict(extra="forbid")
    collection: Optional[str] = Field(None, description="Filter by collection (omit for all).")


class KbRemoveParams(BaseModel):
    model_config = ConfigDict(extra="forbid")
    collection: str = Field(..., description="KB collection name.")
    source: Optional[str] = Field(None, description="Source path to remove (omit to clear entire collection).")


# ---------------------------------------------------------------------------
# KnowledgeTool
# ---------------------------------------------------------------------------


class KnowledgeTool(BaseTool):
    """Agent tool for tenant-scoped knowledge-base ingest and search."""

    @staticmethod
    def _user(execution_context) -> str:
        return getattr(execution_context, "user_id", None) or "_anonymous_"

    @staticmethod
    def _session(execution_context) -> str:
        return getattr(execution_context, "session_id", None) or ""

    @BaseTool.action(
        "Ingest a file or directory into the agent's knowledge base so future kb_search "
        "calls can retrieve it. Skips secrets, binaries, and unchanged files (dedup by hash).",
        param_model=KbIngestParams,
    )
    async def kb_ingest(
        self, params: KbIngestParams, execution_context=None
    ) -> ActionResult:
        user_id = self._user(execution_context)
        session_id = self._session(execution_context)
        try:
            result = await kb_ingest(
                params.path,
                collection=params.collection,
                recursive=params.recursive,
                globs=params.globs,
                user_id=user_id,
                session_id=session_id,
            )
        except Exception as e:
            return ActionResult(error=f"KB ingest failed: {e}", include_in_memory=True)

        if "error" in result:
            return ActionResult(error=result["error"], include_in_memory=True)

        summary = (
            f"KB ingest complete — ingested={result['ingested']}, "
            f"unchanged={result['unchanged']}, chunks={result['n_chunks']}, "
            f"failed={result.get('failed', 0)}, "
            f"skipped(secret={result['skipped_secret']}, "
            f"binary={result['skipped_binary']}, "
            f"office={result.get('skipped_office', 0)}, "
            f"too_large={result.get('skipped_too_large', 0)})"
        )
        return ActionResult(extracted_content=summary, include_in_memory=True)

    @BaseTool.action(
        "Search the agent's knowledge base for relevant content.",
        param_model=KbSearchParams,
    )
    async def kb_search(
        self, params: KbSearchParams, execution_context=None
    ) -> ActionResult:
        from modules.memory.registry import kb_search as _kb_search
        user_id = self._user(execution_context)
        try:
            result = await _kb_search(
                params.query,
                user_id=user_id,
                collection=params.collection,
                limit=params.limit,
            )
        except Exception as e:
            return ActionResult(error=f"KB search failed: {e}", include_in_memory=True)
        if not result:
            return ActionResult(
                extracted_content="No KB results found.", include_in_memory=True
            )
        return ActionResult(extracted_content=result, include_in_memory=True)

    @BaseTool.action(
        "List ingested sources in the knowledge base.",
        param_model=KbListParams,
    )
    async def kb_list(
        self, params: KbListParams, execution_context=None
    ) -> ActionResult:
        from modules.memory.registry import kb_list_sources as _kb_list_sources
        user_id = self._user(execution_context)
        try:
            sources = await _kb_list_sources(
                user_id=user_id, collection=params.collection
            )
        except Exception as e:
            return ActionResult(error=f"KB list failed: {e}", include_in_memory=True)
        if not sources:
            return ActionResult(
                extracted_content="No sources ingested.", include_in_memory=True
            )
        lines = [f"- {s}" for s in sources]
        return ActionResult(
            extracted_content="KB sources:\n" + "\n".join(lines),
            include_in_memory=True,
        )

    @BaseTool.action(
        "Remove a source (or entire collection) from the knowledge base.",
        param_model=KbRemoveParams,
    )
    async def kb_remove(
        self, params: KbRemoveParams, execution_context=None
    ) -> ActionResult:
        from modules.memory.registry import kb_remove as _kb_remove
        user_id = self._user(execution_context)
        try:
            removed = await _kb_remove(
                user_id=user_id,
                collection=params.collection,
                source=params.source,
            )
        except Exception as e:
            return ActionResult(error=f"KB remove failed: {e}", include_in_memory=True)
        return ActionResult(
            extracted_content=f"Removed {removed} chunk(s) from KB.",
            include_in_memory=True,
        )


# ---------------------------------------------------------------------------
# Registration
# ---------------------------------------------------------------------------


def kb_enabled() -> bool:
    from agents.task.constants import AutonomyConfig
    return AutonomyConfig.kb_enabled()


def register_knowledge_tool(force: bool = False) -> bool:
    """Register the 'knowledge' descriptor + class IFF KB_ENABLED (or forced).

    Delegates to ``register_optional_tool``.  No-op when KB is off, so default
    deploys are unaffected.  ``knowledge`` is never in the default ``tool_ids``
    — agents opt in via ``tool_ids=['knowledge']``.
    """
    from tools.descriptors import (
        ToolDescriptor,
        ToolCategory,
        register_optional_tool,
    )

    return register_optional_tool(
        "knowledge",
        KnowledgeTool,
        ToolDescriptor(
            name="knowledge",
            description=(
                "Ingest/search a tenant-scoped knowledge base "
                "(kb_ingest/kb_search/kb_list/kb_remove)"
            ),
            category=ToolCategory.INTEGRATION,
            required_config=[],
            init_priority=80,
            is_optional=True,
        ),
        kb_enabled,
        force=force,
    )
